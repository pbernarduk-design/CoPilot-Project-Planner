import io
from datetime import date
from typing import List, Dict, Any

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# -------------------------------
# Page config
# -------------------------------
st.set_page_config(
    page_title="Project Management Plan Generator",
    page_icon="🗂️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# -------------------------------
# Helper data & utilities
# -------------------------------
METHODOLOGIES = ["Agile", "Waterfall", "Hybrid"]
INDUSTRIES = [
    "Information Technology / Software",
    "Infrastructure / Construction",
    "Healthcare",
    "Financial Services",
    "Government / Public Sector",
    "Education",
    "Other"
]
RISK_APPETITES = ["Low", "Medium", "High"]

DEFAULT_GOV_ROLES = [
    {"Role": "Executive Sponsor", "Name": "", "Decision Rights": "Approves scope, budget, major changes; resolves escalations"},
    {"Role": "Steering Committee", "Name": "", "Decision Rights": "Endorses direction; risk review; executive alignment"},
    {"Role": "Program/Project Director", "Name": "", "Decision Rights": "Governance orchestration; priorities; risk acceptance (within thresholds)"},
    {"Role": "Project Manager", "Name": "", "Decision Rights": "Day-to-day delivery decisions; schedule & risk mgmt; reporting"},
    {"Role": "Product Owner / Business Lead", "Name": "", "Decision Rights": "Backlog priorities; business acceptance"},
    {"Role": "Technical Lead / Architect", "Name": "", "Decision Rights": "Solution design; technical standards"},
    {"Role": "Business Analyst", "Name": "", "Decision Rights": "Requirements approach; traceability"},
    {"Role": "QA / Test Lead", "Name": "", "Decision Rights": "Test strategy; quality gates"},
    {"Role": "Change & Comms Lead", "Name": "", "Decision Rights": "Change impacts; training & comms plan"},
]

# Simple industry & method-aware risk library
BASE_RISKS = {
    "Common": [
        "Scope creep due to unclear requirements or stakeholder changes",
        "Resource constraints or key person dependency",
        "Vendor or third-party delays impacting critical path",
        "Integration complexity leads to unexpected defects",
        "Data privacy/security and compliance gaps",
        "Stakeholder misalignment on priorities and success criteria",
        "Underestimated effort causing schedule slippage",
        "Budget overrun due to change or market costs",
        "Insufficient change management causes poor adoption",
        "Environment or infrastructure readiness delays"
    ],
    "Information Technology / Software": [
        "Requirements churn due to evolving business needs",
        "Technical debt or legacy constraints impacting velocity",
        "Performance/scalability issues discovered late",
        "Environment instability or CI/CD pipeline failures"
    ],
    "Infrastructure / Construction": [
        "Permitting/approvals delays impact start dates",
        "Site conditions differ from surveys impacting design/cost",
        "Weather events disrupt critical path activities",
        "Safety incidents cause stoppages"
    ],
    "Healthcare": [
        "Privacy and clinical safety requirements cause rework",
        "Clinical workflow change resistance limits adoption",
        "Interoperability challenges with EMR/EHR systems"
    ],
    "Financial Services": [
        "Regulatory change introduces additional controls",
        "Audit findings drive unplanned remediation",
        "Compliance approval cycle extends timelines"
    ],
    "Government / Public Sector": [
        "Procurement timelines extend beyond forecast",
        "Policy changes alter project objectives",
        "Public scrutiny/media require additional assurance"
    ],
    "Education": [
        "Academic calendar constraints limit deployment windows",
        "Stakeholder availability limited outside terms",
        "Digital accessibility requirements add scope"
    ]
}


def p_label_to_score(label: str) -> int:
    mapping = {"Low": 1, "Medium": 2, "High": 3}
    return mapping.get(label, 2)


def i_label_to_score(label: str) -> int:
    mapping = {"Low": 1, "Medium": 2, "High": 3}
    return mapping.get(label, 2)


def suggest_objectives(outcome: str, methodology: str, drivers: List[str]) -> List[str]:
    out = []
    if not outcome.strip():
        return out
    out.append(f"Deliver the defined outcome: {outcome.strip()}")

    # SMART-ish prompts
    out.append("Define measurable KPIs and acceptance criteria aligned to business value")
    out.append("Deliver scope in agreed increments with clear Definition of Done" if methodology == "Agile"
               else "Deliver scope by phase with approved entry/exit criteria")
    if "risk" in " ".join(drivers).lower():
        out.append("Reduce high-priority delivery risks with proactive mitigations and fast feedback")
    if "cost" in " ".join(drivers).lower() or "budget" in " ".join(drivers).lower():
        out.append("Control total cost within the approved budget envelope")
    if "speed" in " ".join(drivers).lower() or "time" in " ".join(drivers).lower():
        out.append("Meet or accelerate the target timeline through critical path focus")
    if "quality" in " ".join(drivers).lower():
        out.append("Achieve quality targets measured by defect rates and user satisfaction")
    if "compliance" in " ".join(drivers).lower():
        out.append("Comply with applicable regulatory and policy requirements with evidence")
    return out


def suggest_risks(industry: str, methodology: str) -> List[str]:
    risks = []
    risks.extend(BASE_RISKS["Common"])
    if industry in BASE_RISKS:
        risks.extend(BASE_RISKS[industry])
    if methodology == "Agile":
        risks.append("Backlog not sufficiently refined causing sprint spillover")
        risks.append("Team velocity variability reduces predictability")
    elif methodology == "Waterfall":
        risks.append("Late-stage integration exposes significant defects")
        risks.append("Change control process causes delays on critical path")
    else:  # Hybrid
        risks.append("Hybrid governance causes ambiguity in decision rights")
        risks.append("Phase gates misaligned with agile increments")
    # Deduplicate while preserving order
    seen = set()
    deduped = []
    for r in risks:
        if r not in seen:
            seen.add(r)
            deduped.append(r)
    return deduped


def plan_markdown(model: Dict[str, Any]) -> str:
    """Render final plan to Markdown."""
    lines = []
    L = lines.append

    # Header
    L(f"# Project Management Plan")
    L("")
    L(f"**Project/Outcome:** {model['outcome'] or 'TBD'}")
    L(f"**Methodology:** {model['methodology']}  ")
    L(f"**Industry:** {model['industry']}  ")
    L(f"**Risk Appetite:** {model['risk_appetite']}")
    L("")
    L("---")

    # Objectives
    L("## 1. Objectives")
    if model["objectives"]:
        for i, obj in enumerate(model["objectives"], 1):
            L(f"{i}. {obj}")
    else:
        L("_No objectives captured yet_")
    L("")

    # Scope Summary (lightweight prompt from outcome)
    L("## 2. Scope Summary")
    scope_summary = model.get("scope_summary", "").strip()
    if scope_summary:
        L(scope_summary)
    else:
        L("Summarize in-scope and out-of-scope items, key deliverables, and assumptions.")  # prompt
    L("")

    # Governance
    L("## 3. Governance")
    L(f"**Governance cadence:** {model.get('gov_cadence','TBD')}  ")
    L(f"**Escalation path:** {model.get('gov_escalation','TBD')}")
    L("")
    L("### 3.1 Roles & Decision Rights")
    if model["governance_roles"]:
        df = pd.DataFrame(model["governance_roles"])
        L(df.to_markdown(index=False))
    else:
        L("_No governance roles captured yet_")
    L("")
    if model.get("include_raci"):
        L("### 3.2 RACI (Indicative)")
        L("| Deliverable/Activity | R | A | C | I |")
        L("|---|---|---|---|---|")
        for row in model.get("raci_rows", []):
            L(f"| {row['Item']} | {row['R']} | {row['A']} | {row['C']} | {row['I']} |")
        if not model.get("raci_rows"):
            L("_Add RACI entries in the app to include them here._")
    L("")

    # Risks
    L("## 4. Risks & Mitigations")
    if model["risks"]:
        df = pd.DataFrame(model["risks"])
        df = df.assign(
            ProbabilityScore=df["Probability"].map(p_label_to_score),
            ImpactScore=df["Impact"].map(i_label_to_score),
            RiskScore=lambda d: d["ProbabilityScore"] * d["ImpactScore"]
        ).sort_values("RiskScore", ascending=False)
        # Render selected columns
        L(df[["Risk", "Probability", "Impact", "Mitigation", "Owner"]].to_markdown(index=False))
    else:
        L("_No risks captured yet_")
    L("")

    # Milestones
    L("## 5. Milestones & Timeline")
    if model["milestones"]:
        dfm = pd.DataFrame(model["milestones"])
        L(dfm.to_markdown(index=False))
    else:
        L("_Add key milestones, dates, and acceptance criteria._")
    L("")

    # Stakeholders & Comms
    L("## 6. Stakeholders & Communications")
    if model["comms"]:
        dfc = pd.DataFrame(model["comms"])
        L(dfc.to_markdown(index=False))
    else:
        L("_Define stakeholders, information needs, channels, frequency, and owner._")
    L("")

    # Success Measures
    L("## 7. Success Measures")
    if model["success_measures"]:
        for i, kpi in enumerate(model["success_measures"], 1):
            L(f"{i}. {kpi}")
    else:
        L("_Define measurable KPIs and acceptance criteria._")
    L("")
    return "\n".join(lines)


def plan_docx(markdown_model: Dict[str, Any]) -> bytes:
    """Create a DOCX version of the plan with simple styles."""
    doc = Document()

    # Title
    title = doc.add_paragraph()
    run = title.add_run("Project Management Plan")
    run.bold = True
    run.font.size = Pt(20)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.add_run(f"Project/Outcome: ").bold = True
    meta.add_run(markdown_model['outcome'] or "TBD")
    meta = doc.add_paragraph()
    meta.add_run("Methodology: ").bold = True
    meta.add_run(markdown_model['methodology'])
    meta = doc.add_paragraph()
    meta.add_run("Industry: ").bold = True
    meta.add_run(markdown_model['industry'])
    meta = doc.add_paragraph()
    meta.add_run("Risk Appetite: ").bold = True
    meta.add_run(markdown_model['risk_appetite'])

    def h(txt): 
        p = doc.add_paragraph()
        r = p.add_run(txt)
        r.bold = True
        r.font.size = Pt(14)

    def bullet(txt):
        p = doc.add_paragraph(txt, style="List Bullet")

    doc.add_paragraph()

    # Objectives
    h("1. Objectives")
    if markdown_model["objectives"]:
        for obj in markdown_model["objectives"]:
            bullet(obj)
    else:
        doc.add_paragraph("Define objectives aligned to the outcome and KPIs.")

    # Scope Summary
    h("2. Scope Summary")
    scope_summary = markdown_model.get("scope_summary", "").strip()
    doc.add_paragraph(scope_summary or "Summarize in-scope, out-of-scope, deliverables, and assumptions.")

    # Governance
    h("3. Governance")
    doc.add_paragraph(f"Governance cadence: {markdown_model.get('gov_cadence', 'TBD')}")
    doc.add_paragraph(f"Escalation path: {markdown_model.get('gov_escalation', 'TBD')}")
    doc.add_paragraph("3.1 Roles & Decision Rights").runs[0].bold = True

    if markdown_model["governance_roles"]:
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "Role"
        hdr[1].text = "Name"
        hdr[2].text = "Decision Rights"
        for r in markdown_model["governance_roles"]:
            row = table.add_row().cells
            row[0].text = r.get("Role", "")
            row[1].text = r.get("Name", "")
            row[2].text = r.get("Decision Rights", "")
    else:
        doc.add_paragraph("No governance roles defined.")

    if markdown_model.get("include_raci"):
        doc.add_paragraph("3.2 RACI (Indicative)").runs[0].bold = True
        raci_rows = markdown_model.get("raci_rows", [])
        if raci_rows:
            table = doc.add_table(rows=1, cols=5)
            hdr = table.rows[0].cells
            hdr[0].text = "Deliverable/Activity"
            hdr[1].text = "R"
            hdr[2].text = "A"
            hdr[3].text = "C"
            hdr[4].text = "I"
            for row in raci_rows:
                r = table.add_row().cells
                r[0].text = row.get("Item", "")
                r[1].text = row.get("R", "")
                r[2].text = row.get("A", "")
                r[3].text = row.get("C", "")
                r[4].text = row.get("I", "")
        else:
            doc.add_paragraph("Add RACI entries in the app to include them here.")

    # Risks
    h("4. Risks & Mitigations")
    risks = markdown_model["risks"]
    if risks:
        table = doc.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        hdr[0].text = "Risk"
        hdr[1].text = "Probability"
        hdr[2].text = "Impact"
        hdr[3].text = "Mitigation"
        hdr[4].text = "Owner"
        # sort by score
        df = pd.DataFrame(risks)
        df["ProbabilityScore"] = df["Probability"].map(p_label_to_score)
        df["ImpactScore"] = df["Impact"].map(i_label_to_score)
        df["RiskScore"] = df["ProbabilityScore"] * df["ImpactScore"]
        df = df.sort_values("RiskScore", ascending=False)
        for _, r in df.iterrows():
            row = table.add_row().cells
            row[0].text = str(r["Risk"])
            row[1].text = str(r["Probability"])
            row[2].text = str(r["Impact"])
            row[3].text = str(r.get("Mitigation", ""))
            row[4].text = str(r.get("Owner", ""))
    else:
        doc.add_paragraph("No risks captured yet.")

    # Milestones
    h("5. Milestones & Timeline")
    if markdown_model["milestones"]:
        mt = doc.add_table(rows=1, cols=3)
        mh = mt.rows[0].cells
        mh[0].text = "Milestone"
        mh[1].text = "Target Date"
        mh[2].text = "Acceptance Criteria"
        for m in markdown_model["milestones"]:
            row = mt.add_row().cells
            row[0].text = m.get("Milestone", "")
            row[1].text = str(m.get("Date", ""))
            row[2].text = m.get("Acceptance Criteria", "")
    else:
        doc.add_paragraph("Add key milestones, dates, and acceptance criteria.")

    # Stakeholders & Comms
    h("6. Stakeholders & Communications")
    if markdown_model["comms"]:
        ct = doc.add_table(rows=1, cols=5)
        ch = ct.rows[0].cells
        ch[0].text = "Stakeholder/Group"
        ch[1].text = "Information Needs"
        ch[2].text = "Channel"
        ch[3].text = "Frequency"
        ch[4].text = "Owner"
        for c in markdown_model["comms"]:
            row = ct.add_row().cells
            row[0].text = c.get("Stakeholder", "")
            row[1].text = c.get("Information Needs", "")
            row[2].text = c.get("Channel", "")
            row[3].text = c.get("Frequency", "")
            row[4].text = c.get("Owner", "")
    else:
        doc.add_paragraph("Define stakeholders, information needs, channel, frequency, and owner.")

    # Success measures
    h("7. Success Measures")
    if markdown_model["success_measures"]:
        for k in markdown_model["success_measures"]:
            bullet(k)
    else:
        doc.add_paragraph("Define measurable KPIs and acceptance criteria.")

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def init_state():
    defaults = {
        "outcome": "",
        "methodology": "Hybrid",  # Default methodology (per user preference)
        "industry": "Information Technology / Software",  # Default industry (per user preference)
        "risk_appetite": "Medium",
        "drivers": [],
        "objectives": [],
        "scope_summary": "",
        "governance_roles": DEFAULT_GOV_ROLES.copy(),
        "gov_cadence": "SteerCo monthly; Delivery weekly; Risk review fortnightly",
        "gov_escalation": "PM → Program Director → Sponsor/SteerCo",
        "include_raci": False,
        "raci_rows": [],
        "risks": [],
        "milestones": [],
        "comms": [],
        "success_measures": []
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


init_state()

# -------------------------------
# Sidebar: Configuration
# -------------------------------
with st.sidebar:
    st.header("⚙️ Settings")
    st.session_state["methodology"] = st.selectbox("Delivery methodology", METHODOLOGIES, index=METHODOLOGIES.index(st.session_state["methodology"]))
    st.session_state["industry"] = st.selectbox("Industry", INDUSTRIES, index=INDUSTRIES.index(st.session_state["industry"]))
    st.session_state["risk_appetite"] = st.selectbox("Risk appetite", RISK_APPETITES, index=RISK_APPETITES.index(st.session_state["risk_appetite"]))
    st.markdown("---")
    st.caption("Tip: Settings inform suggestions for objectives and risks.")

st.title("🗂️ Project Management Plan Generator")
st.write("Start with your desired **outcome**, then refine **objectives**, **governance**, and **risks**. Finally, export a complete plan.")

tabs = st.tabs([
    "1) Outcome & Drivers",
    "2) Objectives",
    "3) Governance",
    "4) Risks",
    "5) Milestones & Comms",
    "6) Review & Export"
])

# -------------------------------
# Tab 1: Outcome & Drivers
# -------------------------------
with tabs[0]:
    st.subheader("Define the outcome")
    st.session_state["outcome"] = st.text_area(
        "What outcome are you trying to achieve?",
        value=st.session_state["outcome"],
        placeholder="e.g., Launch a new customer self-service portal that reduces call volume by 30% within 6 months.",
        height=120
    )

    st.subheader("Business drivers (select or write your own)")
    drivers = st.multiselect(
        "Pick applicable drivers:",
        options=["Speed/Time-to-Value", "Cost/Budget Control", "Quality/Reliability", "Risk Reduction", "Regulatory/Compliance", "Customer Experience", "Innovation"],
        default=st.session_state.get("drivers", [])
    )
    custom_drivers = st.text_input("Add any additional drivers (comma-separated)", value="")
    driver_list = drivers + [d.strip() for d in custom_drivers.split(",") if d.strip()]
    st.session_state["drivers"] = driver_list

    st.subheader("Scope summary (optional)")
    st.session_state["scope_summary"] = st.text_area(
        "Briefly summarize in-scope, out-of-scope, key deliverables, assumptions",
        value=st.session_state["scope_summary"],
        placeholder="In-scope: … | Out-of-scope: … | Deliverables: … | Assumptions: …",
        height=120
    )

    if st.button("Generate draft objectives from outcome", type="primary"):
        st.session_state["objectives"] = suggest_objectives(
            st.session_state["outcome"],
            st.session_state["methodology"],
            st.session_state["drivers"]
        )
        st.success("Draft objectives generated. Switch to the **Objectives** tab to refine.")


# -------------------------------
# Tab 2: Objectives
# -------------------------------
with tabs[1]:
    st.subheader("Objectives")
    st.caption("Refine to be specific and measurable (KPIs, dates, thresholds).")
    objs = st.session_state["objectives"] or []
    edited = []
    max_rows = max(4, len(objs))
    for i in range(max_rows):
        val = objs[i] if i < len(objs) else ""
        edited.append(st.text_input(f"Objective {i+1}", value=val, key=f"obj_{i}"))
    st.session_state["objectives"] = [o.strip() for o in edited if o.strip()]

    st.subheader("Success measures (KPIs)")
    kpi_new = st.text_input("Add a KPI (press Enter to add)", key="kpi_new")
    if st.session_state.get("kpi_new") and st.session_state["kpi_new"].strip():
        st.session_state["success_measures"].append(st.session_state["kpi_new"].strip())
        st.session_state["kpi_new"] = ""
    if st.session_state["success_measures"]:
        st.write(pd.DataFrame({"KPI": st.session_state["success_measures"]}))


# -------------------------------
# Tab 3: Governance
# -------------------------------
with tabs[2]:
    st.subheader("Governance")
    st.session_state["gov_cadence"] = st.text_input("Governance cadence", st.session_state["gov_cadence"])
    st.session_state["gov_escalation"] = st.text_input("Escalation path", st.session_state["gov_escalation"])

    st.write("### Roles & Decision Rights")
    gov_df = pd.DataFrame(st.session_state["governance_roles"])
    gov_edited = st.data_editor(
        gov_df,
        use_container_width=True,
        num_rows="dynamic",
        key="gov_editor",
        column_config={"Decision Rights": st.column_config.TextColumn(width="large")}
    )
    st.session_state["governance_roles"] = gov_edited.to_dict(orient="records")

    st.write("### RACI (optional)")
    st.session_state["include_raci"] = st.checkbox("Include a RACI table in the plan", value=st.session_state["include_raci"])
    raci_df = pd.DataFrame(st.session_state.get("raci_rows", []) or [{"Item": "", "R": "", "A": "", "C": "", "I": ""}])
    raci_edit = st.data_editor(
        raci_df,
        use_container_width=True,
        num_rows="dynamic",
        key="raci_editor"
    )
    st.session_state["raci_rows"] = raci_edit.to_dict(orient="records")


# -------------------------------
# Tab 4: Risks
# -------------------------------
with tabs[3]:
    st.subheader("Risks")
    st.caption("Start with suggested risks, then assign Probability, Impact, and Mitigation.")
    col_sg, col_btn = st.columns([4, 1])
    with col_sg:
        suggested = suggest_risks(st.session_state["industry"], st.session_state["methodology"])
        st.write("**Suggested risks (click to add):**")
        for r in suggested:
            if st.button(f"➕ {r}", key=f"addrisk_{hash(r)}"):
                st.session_state["risks"].append({
                    "Risk": r,
                    "Probability": "Medium",
                    "Impact": "Medium",
                    "Mitigation": "",
                    "Owner": ""
                })

    with col_btn:
        st.write("")
        if st.button("Clear all risks"):
            st.session_state["risks"] = []

    risk_df = pd.DataFrame(st.session_state["risks"] or [{"Risk":"", "Probability":"Medium", "Impact":"Medium", "Mitigation":"", "Owner":""}])
    risk_edit = st.data_editor(
        risk_df,
        use_container_width=True,
        num_rows="dynamic",
        key="risk_editor",
        column_config={
            "Probability": st.column_config.SelectboxColumn(options=["Low","Medium","High"]),
            "Impact": st.column_config.SelectboxColumn(options=["Low","Medium","High"]),
            "Mitigation": st.column_config.TextColumn(width="large"),
        }
    )
    st.session_state["risks"] = risk_edit.to_dict(orient="records")

    # Show risk scoring
    if st.session_state["risks"]:
        df_score = pd.DataFrame(st.session_state["risks"])
        df_score["P"] = df_score["Probability"].map(p_label_to_score)
        df_score["I"] = df_score["Impact"].map(i_label_to_score)
        df_score["Risk Score (P×I)"] = df_score["P"] * df_score["I"]
        st.write("**Risk register (scored):**")
        st.dataframe(df_score[["Risk", "Probability", "Impact", "Risk Score (P×I)", "Mitigation", "Owner"]].sort_values("Risk Score (P×I)", ascending=False), use_container_width=True)


# -------------------------------
# Tab 5: Milestones & Comms
# -------------------------------
with tabs[4]:
    st.subheader("Milestones")
    m_df = pd.DataFrame(st.session_state["milestones"] or [{"Milestone":"", "Date":date.today(), "Acceptance Criteria":""}])
    m_edit = st.data_editor(
        m_df,
        use_container_width=True,
        num_rows="dynamic",
        key="m_editor",
        column_config={
            "Date": st.column_config.DateColumn()
        }
    )
    st.session_state["milestones"] = m_edit.to_dict(orient="records")

    st.subheader("Stakeholders & Communications")
    c_df = pd.DataFrame(st.session_state["comms"] or [{"Stakeholder":"", "Information Needs":"", "Channel":"", "Frequency":"", "Owner":""}])
    c_edit = st.data_editor(
        c_df,
        use_container_width=True,
        num_rows="dynamic",
        key="c_editor",
        column_config={"Information Needs": st.column_config.TextColumn(width="large")}
    )
    st.session_state["comms"] = c_edit.to_dict(orient="records")


# -------------------------------
# Tab 6: Review & Export
# -------------------------------
with tabs[5]:
    st.subheader("Review")
    model = {
        "outcome": st.session_state["outcome"],
        "methodology": st.session_state["methodology"],
        "industry": st.session_state["industry"],
        "risk_appetite": st.session_state["risk_appetite"],
        "drivers": st.session_state["drivers"],
        "objectives": st.session_state["objectives"],
        "scope_summary": st.session_state["scope_summary"],
        "governance_roles": st.session_state["governance_roles"],
        "gov_cadence": st.session_state["gov_cadence"],
        "gov_escalation": st.session_state["gov_escalation"],
        "include_raci": st.session_state["include_raci"],
        "raci_rows": st.session_state["raci_rows"],
        "risks": st.session_state["risks"],
        "milestones": st.session_state["milestones"],
        "comms": st.session_state["comms"],
        "success_measures": st.session_state["success_measures"]
    }

    md = plan_markdown(model)
    st.markdown(md)

    st.subheader("Export")
    md_bytes = md.encode("utf-8")
    st.download_button("⬇️ Download Markdown", data=md_bytes, file_name="Project_Management_Plan.md", mime="text/markdown")

    docx_bytes = plan_docx(model)
    st.download_button("⬇️ Download Word (.docx)", data=docx_bytes, file_name="Project_Management_Plan.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

st.caption("Tip: Save your plan artifacts; this app does not persist data between sessions.")
